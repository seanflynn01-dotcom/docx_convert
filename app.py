from flask import Flask, request, jsonify
from docx import Document
from io import BytesIO
import os

app = Flask(__name__)

@app.route('/parse-docx', methods=['POST'])
def parse_docx():
    try:
        # Get the file from the request
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        # Read the document
        doc = Document(BytesIO(file.read()))
        
        # Extract main document text
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract comments
        comments = []
        # Comments are stored in the document's part relationships
        if hasattr(doc, '_part') and hasattr(doc._part, 'rels'):
            for rel in doc._part.rels.values():
                if "comments" in rel.target_ref:
                    comments_part = rel.target_part
                    # Parse comments XML
                    from lxml import etree
                    root = etree.fromstring(comments_part.blob)
                    
                    # Namespace for Word XML
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    
                    for comment in root.findall('.//w:comment', ns):
                        comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                        author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author')
                        date = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date')
                        
                        # Extract comment text
                        comment_text = []
                        for para in comment.findall('.//w:t', ns):
                            if para.text:
                                comment_text.append(para.text)
                        
                        comments.append({
                            'id': comment_id,
                            'author': author,
                            'date': date,
                            'text': ''.join(comment_text)
                        })
        
        return jsonify({
            'document_text': '\n'.join(full_text),
            'comments': comments,
            'total_paragraphs': len(doc.paragraphs),
            'total_comments': len(comments)
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'})

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
```